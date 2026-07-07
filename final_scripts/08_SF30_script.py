"""
08_SF30_script.py

Reads the overview file (produced by 02_overview_file_script.py) and generates the
SF30 Continuation Pages Word document.

The output document contains two tables:

  * Table A - "Attachment J.1 Changes"  : one row per overview row whose 'PR#'
              does NOT start with 'OS'.
  * Table B - "CLIN/Funding Changes"    : one row per overview row whose 'PR#'
              starts with 'OS'. This section always begins on a new page.

Everything else (headings, intro paragraphs, bullet notes) is static boilerplate
that matches the reference document layout.

Only functional programming is used (no classes are defined here).

Dependencies: pandas, python-docx
    pip install pandas python-docx
"""

import os

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ─── Paths & configuration (centralized in config.py) ─────────────────────────
from config import (
    OUTPUT_DIR,
    OVERVIEW_FILE as overview_file,
    SF30_OUTPUT_FILE as sf30_output_file,
    MOD_NUMBER as mod_number,
)


# ─── Constants ────────────────────────────────────────────────────────────────
HEADER_FILL = "4472C4"          # blue shading on header cells
HEADER_FONT_PT = 14
BODY_FONT_PT = 11
FONT_NAME = "Calibri"

TABLE_A_HEADERS = ["PR", "Service", "J.1 Changes"]
TABLE_B_HEADERS = ["Requisition Number", "OpDiv", "Funding Description"]

TABLE_A_WIDTHS = (1.53, 0.86, 4.36)     # inches, per reference document
TABLE_B_WIDTHS = (1.55, 0.74, 4.46)     # inches, per reference document

ROW_HEIGHT_TWIPS = 537                  # minimum row height, per reference document


# ============================================================
# VALUE COERCION HELPERS
# ============================================================

def safe_str(val):
    """Convert *val* to a display string; NaN / None -> empty string."""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return ""
    if isinstance(val, float) and val == int(val):
        return str(int(val))
    return str(val).strip()


def format_version(val):
    """Format the Version value as a zero-padded 'vNN' string (e.g. 1 -> 'v01')."""
    s = safe_str(val)
    if s == "":
        return "v"
    try:
        return f"v{int(float(s)):02d}"
    except (ValueError, TypeError):
        return f"v{s}"


def is_os_pr(pr_value):
    """True iff the PR# starts with 'OS' (case-insensitive)."""
    return safe_str(pr_value).upper().startswith("OS")


def make_change_text(ticket, description):
    """Build the 'Ticket: SF30 Description' string used in both tables."""
    return f"{safe_str(ticket)}: {safe_str(description)}"


# ============================================================
# ROW BUILDERS
# ============================================================

def build_table_a_rows(df):
    """Return list of (PR, Service, J.1 Changes) tuples for non-OS PR rows."""
    rows = []
    for _, r in df.iterrows():
        pr = safe_str(r.get("PR#"))
        if pr == "" or is_os_pr(pr):
            continue
        pr_text = f"{pr} {format_version(r.get('Version'))}"
        service = safe_str(r.get("Services"))
        change = make_change_text(r.get("Ticket"), r.get("SF30 Description"))
        rows.append((pr_text, service, change))
    return rows


def build_table_b_rows(df):
    """Return list of (Requisition Number, OpDiv, Funding Description) tuples for OS PR rows."""
    rows = []
    for _, r in df.iterrows():
        pr = safe_str(r.get("PR#"))
        if not is_os_pr(pr):
            continue
        opdiv = safe_str(r.get("OpDiv"))
        funding = make_change_text(r.get("Ticket"), r.get("SF30 Description"))
        rows.append((pr, opdiv, funding))
    return rows


# ============================================================
# DOCX FORMATTING HELPERS
# ============================================================

def shade_cell(cell, fill_hex):
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def write_cell(cell, text, bold=False, size_pt=BODY_FONT_PT, white=False,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    """Replace a cell's content with a single formatted run, vertically centered."""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    for existing in list(para.runs):
        existing._element.getparent().remove(existing._element)
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = para.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def set_column_widths(table, widths):
    """Force fixed column widths (inches) across every row of the table."""
    table.autofit = False
    table.allow_autofit = False
    for col_idx, grid_col in enumerate(table._tbl.tblGrid.findall(qn('w:gridCol'))):
        if col_idx < len(widths):
            grid_col.set(qn('w:w'), str(int(Inches(widths[col_idx]).twips)))
    for row in table.rows:
        for col_idx, width in enumerate(widths):
            row.cells[col_idx].width = Inches(width)


def set_row_height(row, twips):
    """Apply a minimum (at-least) row height so cells get comfortable padding."""
    row.height = Twips(twips)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def add_data_table(doc, headers, rows, widths):
    """Create a bordered table with a shaded header row and the given data rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_row = table.rows[0]
    set_row_height(header_row, ROW_HEIGHT_TWIPS)
    for col_idx, header in enumerate(headers):
        cell = header_row.cells[col_idx]
        shade_cell(cell, HEADER_FILL)
        write_cell(
            cell, header, bold=True, size_pt=HEADER_FONT_PT, white=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row_values in rows:
        row = table.add_row()
        set_row_height(row, ROW_HEIGHT_TWIPS)
        for col_idx, value in enumerate(row_values):
            # PR / Service / Requisition / OpDiv are centered; the description
            # (last) column is left aligned for readability.
            align = (WD_ALIGN_PARAGRAPH.LEFT if col_idx == len(headers) - 1
                     else WD_ALIGN_PARAGRAPH.CENTER)
            write_cell(row.cells[col_idx], value, align=align)

    set_column_widths(table, widths)
    return table


def add_bold_paragraph(doc, text):
    """Add a Normal paragraph whose single run is bold."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    return para


# ============================================================
# DOCUMENT GENERATION
# ============================================================

def build_document(table_a_rows, table_b_rows, output_path):
    """Assemble and save the SF30 continuation-pages Word document."""
    doc = Document()

    # ── Attachment J.1 Changes (Table A) ──
    doc.add_heading("Attachment J.1 Changes:", level=3)
    doc.add_paragraph(
        "The table below indicates the Attachment J.1 changes by PR including "
        "the associated HHS modification tracking number."
    )
    doc.add_paragraph("")
    add_data_table(doc, TABLE_A_HEADERS, table_a_rows, TABLE_A_WIDTHS)

    # ── Additional Attachment J.1 Changes (static boilerplate) ──
    add_bold_paragraph(doc, "Additional Attachment J.1 Changes:")
    doc.add_paragraph(
        "x new NSC\u2019s added to Tab 5_NSC_Address_Crosswalk.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Moved y end-dated items to Tab 2D_Deleted Items.",
        style="List Bullet",
    )

    # ── CLIN/Funding Changes (Table B) - always on a new page ──
    doc.add_page_break()
    add_bold_paragraph(doc, "CLIN/Funding Changes:")
    doc.add_paragraph(
        "The following table represents Funding Only additions and their "
        "associated requisitions."
    )
    doc.add_paragraph("")
    add_data_table(doc, TABLE_B_HEADERS, table_b_rows, TABLE_B_WIDTHS)

    doc.save(output_path)


# ============================================================
# MAIN
# ============================================================

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("=" * 60)
    print("SF30 CONTINUATION PAGES GENERATION")
    print("=" * 60)
    print(f"Mod Number   : {mod_number}")
    print(f"Input file   : {overview_file}")
    print(f"Output file  : {sf30_output_file}")
    print("=" * 60)

    if not overview_file.exists():
        raise FileNotFoundError(f"Overview file not found: {overview_file}")

    df = pd.read_excel(overview_file)
    print(f"\nOverview rows read: {len(df)}")

    table_a_rows = build_table_a_rows(df)
    table_b_rows = build_table_b_rows(df)

    print(f"Table A (J.1 Changes)    : {len(table_a_rows)} row(s)")
    print(f"Table B (CLIN/Funding)   : {len(table_b_rows)} row(s)")

    print(f"\nGenerating Word doc: {sf30_output_file}")
    build_document(table_a_rows, table_b_rows, sf30_output_file)

    print("\n" + "=" * 60)
    print(f"SF30 output created: {sf30_output_file}")
    print("=" * 60)


if __name__ == "__main__":
    main()
